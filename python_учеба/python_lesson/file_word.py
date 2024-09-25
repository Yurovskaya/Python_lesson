from docx import Document
import re
#  def read_docx(docx_path):
#     # Создаем переменную для сохранения текста всех страниц
#     all_text = ""
#
#      # Открываем файл DOCX
#     doc = Document(docx_path)
#
#      # Итерируем по всем элементам в документе
#     for element in doc.element.body:
#         # Проверяем, является ли текущий элемент таблицей
#          if element.tag.endswith('tbl'):
#              # Итерируем по строкам таблицы
#             for row in element:
#                  # Итерируем по ячейкам в строке
#                 for cell in row:
#                     # Извлекаем текст из абзацев внутри ячейки
#                     cell_text = " ".join(paragraph.text for paragraph in cell if paragraph.text is not None)
#                     all_text += cell_text + "\n"
#
#         # Если текущий элемент не является таблицей, извлекаем текст из параграфов
#          elif element.tag.endswith('p'):
#              if element.text is not None:
#                 all_text += element.text + "\n"
#
#     return all_text
# #
#  docx_path = "/home/yurovskaya/Рабочий стол/python_учеба/python_lesson/Autotest.docx"
#  text_from_docx = read_docx(docx_path)
#  print(text_from_docx)





# import docx2txt
# my_text = docx2txt.process("/home/yurovskaya/Рабочий стол/python_учеба/python_lesson/Autotest.docx")
# print(my_text)


# def read_docx(docx_path):
#     # Создаем переменную для сохранения текста всех страниц
#     all_text = ""
#
#     # Открываем файл DOCX
#     doc = Document(docx_path)
#
#     # Итерируем по всем элементам в документе
#     for element in doc.element.body:
#         # Проверяем, является ли текущий элемент таблицей
#         if element.tag.endswith('tbl'):
#             # Итерируем по строкам таблицы
#             for row in element:
#                 # Итерируем по ячейкам в строке
#                 for cell in row:
#                     # Извлекаем текст из ячейки
#                     cell_text = " ".join(paragraph.text for paragraph in cell if paragraph.text is not None)
#                     all_text += cell_text + "\n"
#
#         # Если текущий элемент не является таблицей, извлекаем текст из параграфов
#         elif element.tag.endswith('p'):
#             if element.text is not None:
#                 all_text += element.text + "\n"
#
#     # Заменяем множественные пробелы на один и убираем лишние пробелы в начале и конце
#     all_text = re.sub(r'\s+', ' ', all_text).strip().replace('. ', '.\n')
#
#     # Проверяем, заканчивается ли текст на заданную строку
#     if all_text.endswith("Схема взаимодействия основных участников отсутствует"):
#         print("Текст оканчивается на нужную строку.")
#     else:
#         print("Текст не оканчивается на нужную строку.")
#
#
#
# docx_path = "/home/yurovskaya/Рабочий стол/python_учеба/python_lesson/Autotest.docx"
# text_from_docx = read_docx(docx_path)
# print(text_from_docx)



# понять что в файле среди контента есть изображение

# from docx import Document
# from PIL import Image
# from io import BytesIO
#
# def has_images(docx_path):
#     doc = Document(docx_path)
#
#     for rel in doc.part.rels.values():
#         if "image" in rel.reltype:
#             return True
#
#     return False
#
# docx_path = "/home/yurovskaya/Рабочий стол/python_учеба/python_lesson/Autotest.docx"
#
# if has_images(docx_path):
#     print("Файл содержит изображения.")
# else:
#     print("Файл не содержит изображений.")


#  узнаем что содержится изображение
# from docx import Document
#
# def has_images(docx_path):
#     doc = Document(docx_path)
#     for rel in doc.part.rels.values():
#         if "image" in rel.reltype:
#             return True
#     return False
#
# def find_image_location(docx_path):
#     doc = Document(docx_path)
#     for rel_id, rel in enumerate(doc.part.rels.values()):
#         if "image" in rel.reltype:
#             print(f"Изображение найдено на странице {rel.target_part.partname}")
#
# docx_path = "/home/yurovskaya/Рабочий стол/python_учеба/python_lesson/Autotest.docx"
#
# if has_images(docx_path):
#     print("Файл содержит изображения.")
#     find_image_location(docx_path)
# else:
#     print("Файл не содержит изображений.")

#
# import os
# from docx import Document
#
#
# def save_images_from_docx(docx_path, output_folder):
#     # Проверяем, существует ли указанная папка для сохранения изображений
#     if not os.path.exists(output_folder):
#         os.makedirs(output_folder)
#
#     # Загружаем документ .docx
#     doc = Document(docx_path)
#
#     # Счетчик для нумерации изображений
#     image_count = 1
#
#     # Переменная для отслеживания наличия изображений
#     images_found = False
#
#     # Перебираем все отношения в документе
#     for rel_id, rel in enumerate(doc.part.rels.values()):
#         # Проверяем, является ли отношение изображением
#         if "image" in rel.reltype:
#             # Получаем данные изображения в виде бинарных данных (blob)
#             image_data = rel.target_part.blob
#
#             # Формируем имя файла для изображения
#             image_filename = f"image_{image_count}.png"
#
#             # Формируем полный путь к файлу изображения
#             image_path = os.path.join(output_folder, image_filename)
#
#             # Записываем бинарные данные изображения в файл
#             with open(image_path, "wb") as image_file:
#                 image_file.write(image_data)
#
#             # Выводим информацию о сохраненном изображении
#             print(f"Saved image {image_count} to: {image_path}")
#
#             # Увеличиваем счетчик для следующего изображения
#             image_count += 1
#
#             # Устанавливаем флаг, что изображение было найдено
#             images_found = True
#
#     # Проверяем, были ли найдены изображения
#     if images_found:
#         print("Изображения в файле существуют.")
#     else:
#         print("В файле не содержится изображений.")
#
#
# # Пример использования
# docx_path = "Autotest.docx"
# output_folder = "папка_для_изображений"
#
# # Вызываем функцию для сохранения изображений из документа .docx
# save_images_from_docx(docx_path, output_folder)

import os
from docx import Document as DocxDocument

def save_images_from_docx(docx_path, output_folder):
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
    doc = DocxDocument(docx_path)
    image_count = 1
    images_found = False
    # Перебираем все отношения в документе
    for rel_id, rel in enumerate(doc.part.rels.values()):
        # Проверяем, является ли отношение изображением
        if "image" in rel.reltype:
            image_data = rel.target_part.blob
            image_filename = f"image_{image_count}.png"
            image_path = os.path.join(output_folder, image_filename)
            with open(image_path, "wb") as image_file:
                image_file.write(image_data)
            print(f"Saved image {image_count} to: {image_path}")
            image_count += 1
            images_found = True
    # Проверяем, были ли найдены изображения
    if images_found:
        print("Изображения в файле существуют.")
    else:
        print("В файле не содержится изображений.")

docx_path = "Autotest.docx"
output_folder = "папка_для_изображений"
save_images_from_docx(docx_path, output_folder)